import os
from dotenv import load_dotenv
from llama_index.core import VectorStoreIndex, StorageContext, Settings, Document
from llama_index.llms.groq import Groq
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
import chromadb
from docx import Document as DocxDocument

load_dotenv()

Settings.llm = Groq(model="llama-3.3-70b-versatile", api_key=os.getenv("GROQ_API_KEY"))
Settings.embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en-v1.5")

DB_PATH = "./data/chroma_db"
COLLECTION_NAME = "documentos_usuario"

def leer_docx(ruta: str) -> str:
    doc = DocxDocument(ruta)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def procesar_y_almacenar(directorio_archivos: str):
    documentos = []
    for archivo in os.listdir(directorio_archivos):
        ruta = os.path.join(directorio_archivos, archivo)
        if archivo.endswith(".docx"):
            texto = leer_docx(ruta)
        else:
            continue  # ignorar otros formatos por ahora
        documentos.append(Document(text=texto, metadata={"archivo": archivo}))

    db = chromadb.PersistentClient(path=DB_PATH)
    chroma_collection = db.get_or_create_collection(COLLECTION_NAME)
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    index = VectorStoreIndex.from_documents(documentos, storage_context=storage_context)

    return f"Procesados {len(documentos)} documentos con éxito."